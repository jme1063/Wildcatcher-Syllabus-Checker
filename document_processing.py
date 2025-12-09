"""
Document Processing Module
This module handles extraction of text from various document formats (PDF, DOCX).
It provides functions to extract text and tables from documents for further processing.
"""

import logging
import pdfplumber
from docx import Document as DocxDocument

# Alternative PDF extraction methods
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False


def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a PDF file using pdfplumber with detailed diagnostics.
    
    Args:
        pdf_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text or None if extraction fails
    """
    text = []
    try:
        with pdfplumber.open(pdf_path) as pdf_doc:
            total_pages = len(pdf_doc.pages)
            logging.info(f"PDF has {total_pages} pages")
            
            for i, page in enumerate(pdf_doc.pages, 1):
                page_text = page.extract_text()
                
                if page_text:
                    page_char_count = len(page_text)
                    text.append(page_text)
                    logging.info(f"Page {i}: Extracted {page_char_count} characters")
                else:
                    logging.warning(f"Page {i}: No text extracted - possible image/scanned page")
                    
                    # Try alternative extraction for image-based pages
                    try:
                        # Check if page has extractable content
                        chars = page.chars
                        if not chars:
                            logging.warning(f"Page {i}: No character objects found - likely image-based")
                        else:
                            logging.info(f"Page {i}: Found {len(chars)} character objects but extract_text() returned nothing")
                    except Exception as e:
                        logging.error(f"Page {i}: Error checking character objects: {e}")
                
                # Extract tables with more detailed logging
                tables = page.extract_tables()
                if tables:
                    logging.info(f"Page {i}: Found {len(tables)} table(s)")
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_text = []
                            for row in table:
                                if row:
                                    table_text.append(" | ".join([str(cell) if cell else "" for cell in row]))
                            if table_text:
                                table_content = "\n".join(table_text)
                                text.append(table_content)
                                logging.info(f"Page {i}, Table {table_idx + 1}: Extracted {len(table_content)} characters")
                
    except Exception as e:
        logging.error(f"Error extracting PDF {pdf_path}: {e}")
        return None
    
    combined_text = "\n".join(text)
    
    if not combined_text.strip():
        logging.warning(f"No text extracted from {pdf_path}")
        return None
    else:
        logging.info(f"TOTAL: Extracted {len(combined_text)} characters from {total_pages} pages in {pdf_path}")
        # Calculate average per page for comparison
        avg_per_page = len(combined_text) / total_pages if total_pages > 0 else 0
        logging.info(f"Average {avg_per_page:.0f} characters per page")
        
        # Flag potentially low extraction and try alternatives
        if avg_per_page < 1000:  # Less than ~1000 chars per page might indicate issues
            logging.warning(f"Low character count per page ({avg_per_page:.0f}) - possible scanned/image-based PDF")
            logging.info("Attempting alternative extraction methods...")
            
            # Try PyPDF2 as alternative
            alternative_text = try_alternative_pdf_extraction(pdf_path)
            if alternative_text and len(alternative_text) > len(combined_text):
                logging.info(f"Alternative extraction yielded {len(alternative_text)} characters (vs {len(combined_text)})")
                return alternative_text
        
        return combined_text


def try_alternative_pdf_extraction(pdf_path):
    """
    Try alternative PDF extraction methods if pdfplumber yields low results.
    
    Args:
        pdf_path (str): Path to the PDF file
        
    Returns:
        str or None: Extracted text from alternative methods
    """
    alternative_text = None
    
    # Try PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pages_text = []
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            pages_text.append(page_text)
                            logging.info(f"PyPDF2 - Page {page_num}: {len(page_text)} characters")
                        else:
                            logging.warning(f"PyPDF2 - Page {page_num}: No text extracted")
                    except Exception as e:
                        logging.error(f"PyPDF2 - Page {page_num} error: {e}")
                
                if pages_text:
                    alternative_text = "\n".join(pages_text)
                    logging.info(f"PyPDF2 total: {len(alternative_text)} characters")
                    
        except Exception as e:
            logging.error(f"PyPDF2 extraction failed: {e}")
    
    # Try PyMuPDF if PyPDF2 didn't work or isn't available
    if (not alternative_text or len(alternative_text) < 5000) and PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(pdf_path)
            pages_text = []
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                if page_text:
                    pages_text.append(page_text)
                    logging.info(f"PyMuPDF - Page {page_num + 1}: {len(page_text)} characters")
                else:
                    logging.warning(f"PyMuPDF - Page {page_num + 1}: No text extracted")
            
            doc.close()
            
            if pages_text:
                pymupdf_text = "\n".join(pages_text)
                logging.info(f"PyMuPDF total: {len(pymupdf_text)} characters")
                
                if not alternative_text or len(pymupdf_text) > len(alternative_text):
                    alternative_text = pymupdf_text
                    
        except Exception as e:
            logging.error(f"PyMuPDF extraction failed: {e}")
    
    return alternative_text


def extract_text_from_docx(docx_path):
    """
    Extracts text from a DOCX file using python-docx with detailed diagnostics.
    
    Args:
        docx_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text or None if extraction fails
    """
    full_text = []
    paragraph_count = 0
    table_count = 0
    header_footer_count = 0
    
    try:
        doc = DocxDocument(docx_path)
        
        # Extract paragraph text with counting
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
                paragraph_count += 1
        
        logging.info(f"Extracted {paragraph_count} paragraphs")
        
        # Extract table content with counting
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    table_text.append(" | ".join(row_data))
            
            if table_text:
                full_text.extend(table_text)
                table_count += 1
                logging.info(f"Table {table_idx + 1}: Extracted {len(table_text)} rows, {sum(len(row) for row in table_text)} characters")
        
        if table_count > 0:
            logging.info(f"Extracted {table_count} tables total")
        
        # Extract from headers and footers if present
        for section_idx, section in enumerate(doc.sections):
            try:
                if section.header.paragraphs:
                    header_text = section.header.paragraphs[0].text.strip()
                    if header_text:
                        full_text.insert(0, header_text)
                        header_footer_count += 1
                        logging.info(f"Section {section_idx + 1} header: {len(header_text)} characters")
                
                if section.footer.paragraphs:
                    footer_text = section.footer.paragraphs[0].text.strip()
                    if footer_text:
                        full_text.append(footer_text)
                        header_footer_count += 1
                        logging.info(f"Section {section_idx + 1} footer: {len(footer_text)} characters")
            except Exception as e:
                logging.warning(f"Could not extract header/footer from section {section_idx + 1}: {e}")
                
    except Exception as e:
        logging.error(f"Error extracting DOCX {docx_path}: {e}")
        return None
    
    combined_text = "\n".join(full_text)
    
    if not combined_text.strip():
        logging.warning(f"No text extracted from {docx_path}")
        return None
    else:
        logging.info(f"TOTAL DOCX EXTRACTION:")
        logging.info(f"  - {paragraph_count} paragraphs")
        logging.info(f"  - {table_count} tables")
        logging.info(f"  - {header_footer_count} headers/footers")
        logging.info(f"  - {len(combined_text)} total characters from {docx_path}")
        
        # Calculate rough page estimate (assuming ~500 words per page, ~5 chars per word)
        estimated_pages = len(combined_text) / 2500
        logging.info(f"  - Estimated ~{estimated_pages:.1f} pages of content")
        
        return combined_text